from tkinter import *
from tkinter import ttk, StringVar, messagebox
from tkinter import END
import webbrowser
import pymysql
from datetime import datetime
import customtkinter as ctk
from tkinter import filedialog
import openpyxl
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import requests
from tkinter import Frame
from docx.shared import RGBColor
import ttkbootstrap  as cttk
from ttkbootstrap.toast import ToastNotification
from ttkbootstrap.tooltip import ToolTip
from ttkbootstrap.constants import *

#-----------------------------------------------


mainFont = font='Heebo'

#----------- database conniction -------------
hostname1 = 'localhost'
porta = 10011
username1 ='root'
passwd1 = 'root'
database1='local'


#--------- | colors | --------------
primary = '#4582ec'
secondary = '#adb5bd'
success = '#02b875'
info = '#17a2b8'
warning = '#f0ad4e'
danger = '#d9534f'
light = '#f8f9fa'
dark = '#343a40'

#--------- | translate | --------------
trns3 = {
    'שם עובד' : 'workername',
    'ת.ז' : 'id',
    'מתאריך' : 'fromdate',
    'עד תאריך' : 'todate',
    'תאריך קבלה' : 'takendate',
}

#-----------------------------------------------
#
#================================================================| |  ניהול עובדים  | |========================================================

class MainScreen:
    def __init__(self):
        super().__init__()
        self.new_root = ctk.CTk()
        self.new_root.geometry('996x780+320+20')
        self.new_root.title('ניהול עוהדים, Develobed by °Anis Zkaria Mhamid°')
        self.new_root.config(bd=0)
        self.new_root.iconbitmap('C:\\NiholOvdem-version-2.5.4\\safety.ico')
        self.new_root.minsize(True,True)
        self.new_root.configure(fg_color=primary)



        def fullscreen(event):
            self.new_root.attributes('-fullscreen', True)

            def exit_fullscreen(event):
                self.new_root.attributes('-fullscreen', False)
            
            self.new_root.bind('<Escape>', exit_fullscreen)
        self.new_root.bind('<F11>', fullscreen)

#=============================================================| |  variables   | |========================================================
        WORKERS_NAME_var = StringVar()
        WORKERS_PHONE_var = StringVar()
        WORKERS_ID_var = StringVar()
        WORKERS_COMPANY_NAME_var = StringVar()
        WORKERS_COMPANY_MANAGER_NAME_var = StringVar()
        WORKERS_WORK_ADDRESS_var = StringVar()
        WORKERS_DATE_var = StringVar()
        WORKERS_WAGE_var = StringVar()
        WORKERS_Taken_var = StringVar()
        WORKERS_Hours_Var = StringVar()
        search_var = StringVar()
        delete_var = StringVar()
        Row_Number_var = IntVar()
# #=====================================


#-----------== | פונקצית שקיפות |
        def allpha(value):
            global alpha
            alpha = float(value)
            alpha_list = [alpha]
            if alpha_list is not None:
                self.new_root.attributes('-alpha', alpha_list)
            else:
                pass
#----------------------------------------


# ============= | |  Header   | | ===================================

        Header = cttk.Frame(self.new_root,borderwidth=0,height=29,bootstyle="primary")
        Header.pack(fill='x',pady=4)
        Header.pack_propagate(False)

        Header3 = cttk.Frame(self.new_root,borderwidth=0,bootstyle="primary")
        Header3.pack(fill='both')

        # main label | ======
        def laBel(master,text,background,foreground,side='top'):
            welcome_to_my_app =cttk.Label(master=master,text=text,background=background, foreground=foreground,font=(mainFont,16,'bold'))
            welcome_to_my_app.pack(side=side)
        laBel(Header,'ניהול עובדים' ,primary,'white')

        # main Button  | | ==========
        def kaftor(text,command,bgcolor,master,side='top'):
            buttons = ctk.CTkButton(master=master,border_color='black',text_color='black',font=(mainFont,17,'bold'),hover_color=success,corner_radius=0,fg_color=bgcolor,text=text,command=command,cursor='hand2')
            buttons.pack(side=side,pady=5,ipady=8,ipadx=20)

#=====================================================================================================================




#================================== | |  תפריט   | | ===================================

# | תמונות תפריט |
#----------------------------------------------------

# | תפריט |
        def toggle_menu():
            point = 0.900

            def toggle_menu_dis():
                tuggle_btn.configure(bootstyle="dark")
                tuggle_btn.configure(command=toggle_menu)
                tuggle_btn.configure(bootstyle="primary")
                for i in range(200):
                    tuggle_menu_frame.place(x=-i-20, y=90, height=window_height, width=20 )
                    tuggle_menu_frame.update()
                    tuggle_menu_frame.after(int(point))

            tuggle_menu_frame = cttk.Frame(master=self.new_root,bootstyle="primary")
            window_height = self.new_root.winfo_height() - 200
            tuggle_menu_frame.place(x=0, y=90, height=window_height, width=200)
            

            laBel(tuggle_menu_frame,'תפריט',primary,light)

            tuggle_btn.configure(bootstyle="dark")
            tuggle_btn.configure(command=toggle_menu_dis)
            tuggle_btn.configure(bootstyle="primary")

            # | כפתור שקיפות |
            cale_tk = ctk.CTkSlider(tuggle_menu_frame,command=allpha,from_=0,to=1,fg_color=success,bg_color=primary)

            laBel(tuggle_menu_frame,'שקיפות',primary,dark)

            for i in range(200):
                tuggle_menu_frame.place(x=i-200, y=80, height=window_height, width=200)
                tuggle_menu_frame.update()
                tuggle_menu_frame.after(int(point))
            
            cale_tk.pack(ipadx=10)

        # | כפתור תפריט |
        tuggle_btn = cttk.Button(Header,width=5,bootstyle="dark", compound='left', command=toggle_menu)
        tuggle_btn.place(x=0,y=0)


#========================================================================================




#======================= | |  צץ תצוגה   | | =============================


# | פונקצית ייצור צץ להצגת הנתונים |
        def create_treeview(parent_frame):
            treeview_frame = ttk.Frame(parent_frame)
            treeview_frame.pack(padx=0, pady=0, fill=BOTH, expand=True)

            vertical_scrollbar = ttk.Scrollbar(treeview_frame,bootstyle='danger', orient="vertical")
            vertical_scrollbar.pack(side='right', fill='y')

            horizontal_scrollbar = ttk.Scrollbar(treeview_frame,bootstyle='danger', orient="horizontal")
            horizontal_scrollbar.pack(side='bottom', fill='x')


            workers_treeview = ttk.Treeview(
                treeview_frame,
                padding=(0, 0),
                xscrollcommand=horizontal_scrollbar.set,
                yscrollcommand=vertical_scrollbar.set,
                columns=('Rowid', "hours", 'taken', "wage", 'address', 'managername', 'compname', 'date', 'worker_id', 'phone', 'workername')
            )
            workers_treeview.pack(fill='both', expand=True)
            workers_treeview.config(selectmode='extended', height=13)

            horizontal_scrollbar.configure(command=workers_treeview.xview)
            vertical_scrollbar.configure(command=workers_treeview.yview)
            workers_treeview.xview_moveto(1)

            workers_treeview.heading('Rowid', text="מ'ס שורה")
            workers_treeview.heading('hours', text="מ'ס שעות")
            workers_treeview.heading('taken', text='מפריעה')
            workers_treeview.heading('wage', text='שכר')
            workers_treeview.heading('address', text='כתובת')
            workers_treeview.heading('managername', text='מנהל')
            workers_treeview.heading('compname', text='חברה')
            workers_treeview.heading('date', text='תאריך')
            workers_treeview.heading('worker_id', text="ת'ז")
            workers_treeview.heading('phone', text='טלפון')
            workers_treeview.heading('workername', text='שם')

            workers_treeview.column('#0', width=0, stretch=NO)
            workers_treeview.column('Rowid', width=40)
            workers_treeview.column('hours', width=40)
            workers_treeview.column('taken', width=50)
            workers_treeview.column('wage', width=50)
            workers_treeview.column('address', width=60)
            workers_treeview.column('managername', width=80)
            workers_treeview.column('compname', width=70)
            workers_treeview.column('date', width=70)
            workers_treeview.column('worker_id', width=80)
            workers_treeview.column('phone', width=80)
            workers_treeview.column('workername', width=70)

            return workers_treeview
        workers_tuple = create_treeview(self.new_root)

#================ | |  Functions   | | ================================
        
# | פונקצית תוסט להוצאת השגיאות למשתמש 
        def toastErrorCacher(title,message):
            toast = ToastNotification(title=title,message=message,duration=10000,alert=True)
            toast.show_toast()

# | פונקצית הסתרת צץ הנתונים |
        def toggle_treeview():
                if workers_tuple.winfo_ismapped():
                    workers_tuple.pack_forget()
                else:
                    workers_tuple.pack(fill='both')


# | פונקצית קבלת הנתונים |==
        def get_api():
            try:
                # GET request to the API endpoint
                response = requests.get('http://127.0.0.1:5000/get')
                
                # Verify the success of the request
                if response.status_code == 200:
                    # Extract data into JSON format
                    data = response.json()
                    
                    # Clear existing treeview data
                    workers_tuple.delete(*workers_tuple.get_children())
                    
                    # Inserting new data from the API into the Treeview
                    for row in data:
                        workers_tuple.insert('', END, values=(
                            row['Rowid'],
                            row['hours'],
                            row['taken'],
                            row['wage'],
                            row['address'],
                            row['managername'],
                            row['compname'],
                            row['date'],
                            row['worker_id'],
                            row['phone'],
                            row['workername']
                        ))
                else:
                    # If the request fails
                    messagebox.showerror('Error', f'Failed to fetch data. Status code: {response.status_code}')
            except requests.exceptions.RequestException as e:
                # If any exceptions occur during the order
                messagebox.showerror('Error', f'An error occurred: {e}')
        get_api()


# | פונקצית שליחת הנתונים | 
        def post_api():
            # Assuming  necessary data stored in variables
            data = {
                'WORKERS_NAME_var': WORKERS_NAME_var.get(),
                'WORKERS_PHONE_var': WORKERS_PHONE_var.get(),
                'WORKERS_ID_var': WORKERS_ID_var.get(),
                'WORKERS_COMPANY_NAME_var': WORKERS_COMPANY_NAME_var.get(),
                'WORKERS_COMPANY_MANAGER_NAME_var': WORKERS_COMPANY_MANAGER_NAME_var.get(),
                'WORKERS_WORK_ADDRESS_var': WORKERS_WORK_ADDRESS_var.get(),
                # 'WORKERS_DATE_var': WORKERS_DATE_var.get(),
                'WORKERS_WAGE_var': WORKERS_WAGE_var.get(),
                'WORKERS_Taken_var': WORKERS_Taken_var.get(),
                'WORKERS_Hours_Var': WORKERS_Hours_Var.get(),
            }

            # Set mainContent-Type header to application/json
            headers = {'mainContent-Type': 'application/json'}

            # تحديد عنوان URL للطلب
            url = 'http://127.0.0.1:5000/post'

            try:
                response = requests.post(url, json=data, headers=headers)

                if response.status_code == 201:
                    # عرض رسالة نجاح إذا كان الرد 201 Created
                    messagebox.showinfo('Success', 'Data added successfully')
                else:
                    # التعامل مع أخطاء محددة إذا فشلت عملية الإضافة
                    if response.status_code == 400:
                        messagebox.showerror('Error', 'Bad request. Please check your data.')
                    elif response.status_code == 500:
                        messagebox.showerror('Error', 'Internal server error. Please try again later.')
                    else:
                        messagebox.showerror('Error', f'Failed to add data. Status code: {response.status_code}')

            except requests.exceptions.RequestException as e:
                # التعامل مع الأخطاء العامة في الطلبات (مثل مشاكل الشبكة)
                messagebox.showerror('Error', f'An error occurred: {e}')
            except Exception as e:
                # التعامل مع أي استثناءات غير متوقعة
                messagebox.showerror('Error', f'An unexpected error occurred: {e}')

# | פונקצית עדכון הנתונים |==
        def update_api():
            # data stored in variables
            data = {
                'Rowid': Row_Number_var.get(),
                'hours': WORKERS_Hours_Var.get(),
                'taken': WORKERS_Taken_var.get(),
                'wage': WORKERS_WAGE_var.get(),
                'address': WORKERS_WORK_ADDRESS_var.get(),
                'managername': WORKERS_COMPANY_MANAGER_NAME_var.get(),
                'compname': WORKERS_COMPANY_NAME_var.get(),
                'date': WORKERS_DATE_var.get(),
                'worker_id': WORKERS_ID_var.get(),
                'phone': WORKERS_PHONE_var.get(),
                'workername': WORKERS_NAME_var.get(),
            }
            # Set mainContent-Type header to application/json
            headers = {'mainContent-Type': 'application/json'}

            # API endpoint URL
            url = 'http://127.0.0.1:5000/update'
            

            try:
                # Send PUT request to the API endpoint
                response = requests.put(url, json=data,headers=headers)
                if response.status_code == 200:
                    messagebox.showinfo('Success', 'Data updated successfully')
                else:
                    print(response)
                    # Handle specific error when updating data fails
                    if response.status_code == 400:
                        messagebox.showerror('Error', 'Bad request. Please check your data.')
                    elif response.status_code == 500:
                        messagebox.showerror('Error', 'Internal server error. Please try again later.')
                    else:
                        messagebox.showerror('Error', f'Failed to update data. Status code: {response.status_code}')
            except requests.exceptions.RequestException as e:
                # Handle general request exceptions (e.g., network issues)
                messagebox.showerror('Error', f'An error occurred: {e}')
            except Exception as e:
                # Handle any other unexpected exceptions
                messagebox.showerror('Error', f'An unexpected error occurred: {e}')


# | פונקצית מחיקת נתונים |==
        def delete_worker():
            try:
                # استدعاء الدالة delete_api مع رقم السجل المراد حذفه في عنوان URL
                worker_id = delete_var.get()
                response = requests.delete(f'http://127.0.0.1:5000/delete/{worker_id}')
                if response.status_code == 200:
                    messagebox.showinfo('Success', 'Data deleted successfully')
                elif response.status_code == 404:
                    messagebox.showerror('Error', 'Worker not found')
                else:
                    messagebox.showerror('Error', f'Failed to delete data. Status code: {response.status_code}')
            except requests.exceptions.RequestException as e:
                # Handle general request exceptions (e.g., network issues)
                messagebox.showerror('Error', f'An error occurred: {e}')
            except Exception as e:
                # Handle any other unexpected exceptions
                messagebox.showerror('Error', f'An unexpected error occurred: {e}')


# | פונקצית ניכוי כל שדות כל  |==
        def Clear_function():

            WORKERS_Hours_Var.set('')

            WORKERS_Taken_var.set('')
            
            WORKERS_WAGE_var.set('')

            WORKERS_WORK_ADDRESS_var.set('')

            WORKERS_COMPANY_MANAGER_NAME_var.set('')

            WORKERS_COMPANY_NAME_var.set('')

            WORKERS_ID_var.set('')

            WORKERS_PHONE_var.set('')

            WORKERS_NAME_var.set('')

            delete_var.set('')

            search_var.set('')


# | פונקצית קבלה | ניקוי שדות כל  |==
        def addandclear_function():
            post_api()
            Clear_function()
            get_api()
            get_workers_names()
            get_workers_names2(search2=search2)
            get_workers_names_Analysis()
            workers_tuple.xview_moveto(1)


# | פונקצית מחיקה | קבלה | ניקוי | שדות כל  |==
        def DeleteAndClear_function():
            delete_worker()
            get_api()
            get_workers_names()
            get_workers_names_Analysis()
            get_workers_names2(search2=search2)
            Clear_function()


# | ממלא את שדות הקלט כאשר אתה לוחץ על שם |
        def get_cursor_function(event):
            cursor_row = workers_tuple.focus()
            try:    
                if cursor_row:
                    mainContents = workers_tuple.item(cursor_row)
                    row = mainContents['values']
                    if row:
                        Row_Number_var.set(row[0])
                        WORKERS_Hours_Var.set(row[1])
                        WORKERS_Taken_var.set(row[2]),
                        WORKERS_WAGE_var.set(row[3])
                        WORKERS_WORK_ADDRESS_var.set(row[4])
                        WORKERS_COMPANY_MANAGER_NAME_var.set(row[5])
                        WORKERS_COMPANY_NAME_var.set(row[6])
                        WORKERS_DATE_var.set(row[7])
                        WORKERS_ID_var.set(row[8])
                        WORKERS_PHONE_var.set(row[9])
                        WORKERS_NAME_var.set(row[10])
            except pymysql.err.IntegrityError as e:
                messagebox.showerror('err', e)
            finally:
                pass
        workers_tuple.bind("<ButtonRelease-1>", get_cursor_function)


#| מחקה באמצעות כפתור delete |
        def delete_row_function(event):
            try:
                con = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                with con.cursor() as cor:
                    selected_items = workers_tuple.selection()
                    if selected_items:
                        for selected_item in selected_items:
                            row_id = workers_tuple.item(selected_item, 'values')[0]
                            confirmation = messagebox.askyesno("מוחק שורה", f"?האם אתה בטוח שברצונך למחוק שורה:_{row_id}")
                            if confirmation:
                                delete_query = "DELETE FROM workers WHERE `Rowid` = %s"
                                cor.execute(delete_query, (row_id,))
                                workers_tuple.delete(selected_item)
                            else:
                                pass
                        con.commit()
                        get_workers_names()
                        get_workers_names2(search2=search2)
                        dataAnalysis_func()
                    else:
                        toastErrorCacher("שגיאה", "לא נמחק")
            finally:
                pass
        workers_tuple.bind('<Delete>', delete_row_function)
        workers_tuple.bind('<Return>', post_api) 


        def validate_number(En1) -> bool:
            """Validates that the input is a number"""
            if En1.isdigit():
                return True
            elif En1 == "":
                return True
            else:
                return False

        def validate_alpha(En1) -> bool:
            """Validates that the input is alpha"""
            if En1.isdigit():
                return False
            elif En1 == "":
                return True
            else:
                return True


# | פונקצית עדכון נתונים במסד הנתונים | ניכוי שדות כלט  |
        def updateandclear():
            update_api()
            get_api()
            Clear_function()
            dataAnalysis_func()
            get_workers_names2(search2)
            get_workers_names_Analysis()
            workers_tuple.xview_moveto(1)


        def Search_function(event=None):
            try:
                searcha = search.get()
                con = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                cor = con.cursor()

                cor.execute("SELECT * FROM workers WHERE workername = %s", (searcha,))
                quary = cor.fetchall()

                if quary:
                    workers_tuple.delete(*workers_tuple.get_children())
                    for row in quary:
                        workers_tuple.insert('', END, values=row)
                    workers_tuple.xview_moveto(1.0)
                    con.commit()
                else:
                    toastErrorCacher("החיפוש שלך לא תאם","לא נמצאו עובדים התואמים לקריטריוני החיפוש שלך")
            except pymysql.Error as e:
                toastErrorCacher("שגיאה", "אירעה שגיאה בעת בדיקת מסד הנתונים")
            finally:
                con.close()


        #=== | פונקצית חישוב שכר | 
        def calculate_sum_function():
            try:
                conn = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                cursor = conn.cursor()
                sby = search2.get()
                if sby == 'כל העובדים':
                    query = f"SELECT SUM(`wage`) FROM workers"
                    cursor.execute(query)
                    res = cursor.fetchone()
                    total_wage = res[0]
                    toastErrorCacher('חישוב סך הכל',f'הסכום הכולל של כל העובדים: {total_wage}')
                    return total_wage
                else:
                    query = f"SELECT workername, SUM(`wage`) AS total_wage, SUM(Taken) AS total_taken FROM workers WHERE workername = '{sby}' GROUP BY workername"
                    cursor.execute(query)
                    res = cursor.fetchall()
                    if res:
                        for row in res:
                            workername, total_wage, total_taken = row
                            Min = total_wage - total_taken
                            result = messagebox.showinfo('חישוב סך הכל', f'פרוטו עבור {workername} : {total_wage} \n\nקיבל :{total_taken} \n\n נטו : {Min}')
                            result
                            return total_wage, total_taken, Min
                    else:
                        toastErrorCacher("שגיאה","אין תוצאות, בחר שם עובד")
            except Exception as e:
                        toastErrorCacher('לא ניתן להתחבר לשרת MySQL ',f" לא ניתן היה ליצור חיבור מכיוון שמכונת היעד סירבה לכך באופן פעיל בדוק את חיבור האינטרנט שלך ונסה שנית")
            finally:
                return None


#-------------- | תאריך | ------------------ 
        def update_date():
            current_date = datetime.now().strftime('%d/%m/%Y')
            global date_entry
            date_entry = cttk.Entry(Header3, font=(mainFont, 14, 'normal'),justify=CENTER, state='normal')
            date_entry.insert(0, current_date)
            date_entry.pack(side='left',padx=0)
            WORKERS_DATE_var.set(current_date)
        update_date()


#--------------- | פונקצית שעון | ----------------
        def update_time():
            current_time = datetime.now()
            current_hour = current_time.hour
            current_minute = current_time.minute
            current_second = current_time.second
            formatted_time = "{} : {} : {}".format(current_hour, current_minute, current_second)
            date_entry2.delete(0, END)
            date_entry2.insert(0, formatted_time)
            Header3.after(1000, update_time) # |שדה שעון |
        date_entry2 = cttk.Entry(Header3,
                                justify=CENTER,
                                font=(mainFont, 14, 'normal'),
                                state='normal')
        date_entry2.pack(side='right')
        update_time()


        # | פונקצית צור קשר |
        def infoo():
            toastErrorCacher('צור קשר' , 'Email: anesmhamed1@gmail.com\n\nPhone: 0538346915\n')
            url = "https://anismhamid.github.io/anis-mhamid-project/"
            webbrowser.open(url)


        # | kill program פונקצית|
        def des():
            self.new_root.destroy()



#======= | |  Printing Functions  | | =========================


        # Customize the table in DOCX
        def customize_table(table):
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0, 0, 0)  # Black font color
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        customize_table(table)
        doc.save("output.docx")

        # Print to DOCX
        def print_to_docx():
            selected_value = search2.get()
            default_file_name = f'{selected_value}'
            file_path_docx = filedialog.asksaveasfilename(defaultextension=".docx",
                                                        initialfile=default_file_name,
                                                        filetypes=[("Word Files", "*.docx")])
            if file_path_docx:
                document = Document()
                title = document.add_heading(f"\n\n מידע עובדים\n", level=1)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                table_data = []
                column_names = [workers_tuple.heading(column)["text"] for column in workers_tuple["columns"]]
                table_data.append(column_names)
                for item in workers_tuple.get_children():
                    item_values = [workers_tuple.item(item, "values")[i] for i in range(len(workers_tuple["columns"]))]
                    table_data.append(item_values)
                table = document.add_table(rows=len(table_data), cols=len(column_names))
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = str(cell_data)
                customize_table(table)
                document.save(file_path_docx)
                messagebox.showinfo('נשמר בהצלחה', f"{file_path_docx} :הקובץ נשמר ב")


        # Print to Excel
        def print_to_excel():
            selected_value = search.get()
            default_file_name2 = f'{selected_value}'
            try:
                file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                        initialfile=default_file_name2,
                                                        filetypes=[("Excel Files", "*.xlsx")])
                if file_path:
                    workbook = openpyxl.Workbook()
                    sheet = workbook.active
                    sheet.sheet_view.rightToLeft = True
                    columns = workers_tuple["columns"]
                    column_headers = [workers_tuple.heading(column)["text"] for column in columns]
                    column_headers.reverse()
                    sheet.append(column_headers)
                    items = workers_tuple.get_children()
                    for row_idx, item in enumerate(items, start=2):
                        values = [workers_tuple.item(item, "values")[columns.index(column)] for column in columns]
                        values.reverse()
                        for col_idx, value in enumerate(values, start=1):
                            sheet.cell(row=row_idx, column=col_idx, value=value)
                    for col_idx, column_header in enumerate(column_headers, start=1):
                        sheet.cell(row=1, column=col_idx, value=column_header)
                        sheet.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = len(column_header) + 2
                    workbook.save(file_path)
                    if workbook:
                        messagebox.showinfo('הצלחה', f"הקובץ נשמר ב: {file_path}")
            except FileNotFoundError:
                messagebox.showerror('שגיאה', 'הקובץ לא נמצא.')
            except PermissionError:
                messagebox.showerror('שגיאה', 'הקובץ לא ניתן לכתיבה. נא לסגור אותו ולנסות שוב.')
            except Exception as e:
                messagebox.showerror('שגיאה', f'אירעה שגיאה: {str(e)}')



#       
        def mainContentView():
            f8.pack_forget()
            f7.pack(fill='both')

        def tloshimSectionView():
            f7.pack_forget()
            f8.pack(fill='both')
# ==================================================================



#================== | כפתורים עליונים | =================

        def topButtons(text,command,side):
            _top = ctk.CTkButton(Header3,
                                corner_radius=0,
                                width=150,
                                font=(mainFont,15,"bold"),
                                text=text,
                                command=command)
            _top.pack(side=side,padx='20')
        topButtons(text='תלושים',command=tloshimSectionView,side=LEFT)
        topButtons(text='הסתרת רשימה',command=toggle_treeview,side=RIGHT)
        topButtons(text='ראשי',command=mainContentView,side=RIGHT)



#===================================================



#=======| |  Calclate   | |=========

#                         | פרים חישוב | חיצוני  |
        fram_heshov = cttk.Frame(self.new_root,relief='sunken',width=100)
        fram_heshov.pack(anchor='w',fill='x',pady=0,ipady=0)
#----------------------------------------------------

#                         | פרים חישוב פנימי  |
        fram_heshov2 = Frame(fram_heshov,bg='black',relief='sunken',bd=0)
        fram_heshov2.pack(anchor='s',fill='x',ipadx=0)
#----------------------------------------------------

#                        | בחירת שם לחישוב |
        search2 = ttk.Combobox(fram_heshov2, state='readonly',width=15, font=( mainFont,12 ), justify='center')
        search2.pack(side='left',padx=(0,0),pady=0)
#----------------------------------------------------

#                 | פונקצית ייבואי שימות מדאטאביס 2 |
        def get_workers_names2(search2):
            try:
                connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                cursor = connection.cursor()
                cursor.execute("SELECT DISTINCT workername FROM workers ORDER BY workername ASC;")
                rows = cursor.fetchall()
                global worker_names
                worker_names = [row[0] for row in rows]
                search2['values'] =['כל העובדים']+ worker_names
            except:
                pass
        get_workers_names2(search2)
#---------------------------------------------------------


#                         | כפתור חישוב |                
        calcbtn = ctk.CTkButton(master = fram_heshov2,
                                text='חישוב',
                                width=80,
                                fg_color=primary,
                                hover_color=danger,
                                font=(mainFont ,18,'bold'),
                                command=calculate_sum_function)
        calcbtn.pack(side=LEFT,padx=50)


#==========================================================| |  Search  | |===========================================

        refreshFrame = Frame(master=fram_heshov2 ,bg=light,relief="solid",border=0)
        refreshFrame.pack(side='left',padx=(10,0),ipady=0)

        # | כפתור הצג הכל| 
        kaftor(text='רענן טבלה',master=refreshFrame,command=get_api,bgcolor=primary)




# ------------------- | בחירת שם חיפוש | -----------------------------------
        laBel(fram_heshov2,'חיפוש',light,dark)

        search = cttk.Combobox(fram_heshov2,style="Custom.TCombobox", state='readonly',width=15, justify='center')
        search.pack(side='right', padx=(10,10))
        search.bind("<<ComboboxSelected>>", Search_function)


        def get_workers_names():
            try:
                connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                cursor = connection.cursor()

                cursor.execute("SELECT DISTINCT workername FROM workers ORDER BY workername ASC;")

                rows = cursor.fetchall()

                global worker_names
                worker_names = [row[0] for row in rows]
            
                search['values'] = worker_names

            except:
                pass
        get_workers_names()
#=====================================================================================================================




        f7 = Canvas(self.new_root, bg='gray', bd=0, highlightthickness=0)
        f7.pack(pady=2, fill='both', expand=True)


#============================= | |  Main Content Area   | | ====================================

        #------------------- |  מחיקה | ---------

        def mainContent():
            deleteframe = Frame(f7,bg=danger,relief='sunken',bd=5,)
            deleteframe.pack(side=LEFT,fill='y')

            laBel(deleteframe,"מחק לפי מ'ס שורה",light,dark)

            En_Delete = ctk.CTkEntry(deleteframe,
                                    textvariable=delete_var,
                                    justify='center',
                                    font=(mainFont ,14),
                                    border_width=1,
                                    text_color='white',
                                    border_color='red',fg_color='red')
            En_Delete.pack(pady=0)
            En_Delete.bind('<Enter>', lambda e: En_Delete.configure(border_color='#3cff00'))
            En_Delete.bind('<Leave>', lambda e: En_Delete.configure(border_color='red'))


            #==  כפתור מחיקה |=

            kaftor('מחיקה',DeleteAndClear_function,danger,deleteframe,side='bottom')
    #=====================================================================================================================





    #======================== | שדות כלט | =============================
            f1 = cttk.Frame(f7,bootstyle='primary',relief='sunken',border=5)
            f1.pack(side='right',fill='y',pady=0)

            def en_Title():
                global entry_and_entryTitle
                global En1
                def entry_and_entryTitle(master,validcommand, valid,textvar, placehold):
                    En1 = cttk.Entry(master=master,bootstyle='default',takefocus=placehold,validatecommand=validcommand,validate=valid, textvariable=textvar, font=(mainFont, 10,'bold'), justify='center')
                    En1.insert(0, placehold)
                    En1.bind('<Enter>', lambda e: on_entry_focus_in(En1, placehold))
                    En1.bind('<Leave>', lambda e: on_entry_focus_out(En1, placehold))
                    En1.pack(pady=(3,0))

                    def on_entry_focus_in(entry, placeholder):
                        if entry.get() == placeholder:
                            entry.delete(0, 'end')

                    def on_entry_focus_out(entry, placeholder):
                        if not entry.get():
                            entry.insert(0, placeholder)

                # register the validation callback
                global alpha_func
                digit_func = self.new_root.register(validate_number)
                alpha_func = self.new_root.register(validate_alpha)

                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_NAME_var ,placehold = 'שם')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_PHONE_var ,placehold = 'טלפון')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_ID_var ,placehold= 'ת.ז')
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_DATE_var ,placehold = '')
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_COMPANY_NAME_var ,placehold = 'שם חברה')
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_COMPANY_MANAGER_NAME_var ,placehold = ' מנהל עבודה') 
                entry_and_entryTitle(master = f1 , validcommand = (alpha_func, '%P'), valid='focus' , textvar = WORKERS_WORK_ADDRESS_var ,placehold = 'כתובת')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_WAGE_var ,placehold = 'שח')
                entry_and_entryTitle(master = f1 , validcommand = (digit_func, '%P'), valid='focus' , textvar = WORKERS_Taken_var ,placehold = '0')


                #========================= | כפתור הןספה | ======================
                button1 = ctk.CTkButton(f1,corner_radius=0,hover_color=danger, text='הוספת', fg_color=success, text_color='black', font=(mainFont, 16, 'bold'), command=addandclear_function, cursor='hand2 ')
                button1.pack(side=BOTTOM,pady=(0,10))
                ToolTip(button1,delay=250, text="מלא את שדות הכלט מטה לפני ההוספה", bootstyle=('dark', INVERSE))
            en_Title()

    #=======================================================================



# ===================== | |  Buttons frame | | =======================================

# ===================== | פרים כפתורים | =======================================
            bottonFrame = Frame(f7,bd=3,relief='flat',bg='gray')
            bottonFrame.pack(side='right',anchor='nw',fill='y')


            entry_and_entryTitle(master=bottonFrame,validcommand=(alpha_func, '%P'),valid='focus',textvar=WORKERS_Hours_Var,placehold=str("מ'ס שעות"))

            kaftor(text='תיקון',command=updateandclear,bgcolor=primary,master=bottonFrame)

            kaftor(text="Docx הדפסה ל",command=print_to_docx,bgcolor=warning,master=bottonFrame)

            kaftor(text="Exel הדפסה ל",command=print_to_excel,bgcolor=warning,master=bottonFrame)

            kaftor(text='צור קשר',command=infoo,bgcolor=info,master=bottonFrame)

            kaftor(text='נקה שדות כלט',command=Clear_function,bgcolor=warning,master=bottonFrame)

            kaftor(text='סגור תוכנה',command=des,bgcolor=danger,master=bottonFrame)
#======================================================================================


            # | פרים Rowid | -----------------------------
            rowIdFrame = Frame(f7,bg='blue',relief='sunken',bd=3)
            rowIdFrame.pack(fill='both',ipady=5)
            #----------------------------------------------------------


            # ====================== | ניתוח נתונים |
            
            # | פרים ניתוח נתונים | -----------------------------
            underprintFrame = Frame(f7,bg=primary,relief='sunken',bd=1)
            underprintFrame.pack(pady=0,fill='both',ipady=300)
            

            global dataAnalysis_func
            def dataAnalysis_func():
                global get_workers_names2
                calc_Days_Frame = cttk.Frame(underprintFrame,bootstyle="primary")
                calc_Days_Frame.pack(fill='both',ipady=300)

                # =========== | לאביל ניתוח נתונים | =============
                laBel(calc_Days_Frame,'ניתוח נתונים',primary,light)

                search_ttk = cttk.Combobox(master=calc_Days_Frame,
                                        style="Custom.TCombobox",
                                        validate='all',
                                        state='readonly',
                                        font=(mainFont,10),
                                        justify='center')
                search_ttk.pack(side='top',pady=(10, 20))

                global get_workers_names_Analysis
                def get_workers_names_Analysis():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        cursor = connection.cursor()

                        cursor.execute("SELECT DISTINCT workername FROM workers;")

                        rows = cursor.fetchall()

                        worker_names = [row[0] for row in rows]

                        search_ttk['values'] = [''] + worker_names

                    except pymysql.err.DatabaseError as e:
                        messagebox.showerror("no Interner Error",f"אירעה שגיאה בעת חיבור למסד הנתונים {e}")
                    finally:
                        pass
                get_workers_names_Analysis()

                def howMutchDays():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        
                        cursor = connection.cursor()
                        cursor.execute(f"SELECT COUNT(*) AS name_count FROM workers WHERE workername = '{search_ttk.get()}'")
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[0]
                            fresh = '      '
                            messagebox.showinfo("",f"מספר הימים ל {search_ttk.get()}\n{fresh}   {name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR","אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        messagebox.showerror("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                    finally:
                        pass

                def takenMoney():
                    try:

                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        
                        cursor = connection.cursor()

                        cursor.execute(f"SELECT workername, SUM(`Taken`) AS total_wage FROM workers WHERE workername = '{search_ttk.get()}'")
                        
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[1]
                            fresh = '      '
                            messagebox.showinfo("", f"   קיבל מפריעה  \n{fresh}   {name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR", "אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        # Handling database errors
                        messagebox.showerror("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                    finally:
                        pass

                def howMutchDaysAll():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                        cursor = connection.cursor()
                        cursor.execute("SELECT COUNT(DISTINCT workername) AS name_count FROM workers")
                        result = cursor.fetchone()
                        
                        if result:
                            name_count = result[0]
                            messagebox.showinfo("מספר העובדים:",f"מספר העובדים הוא: {search_ttk.get()}{name_count} ")
                                
                        else:
                            messagebox.showinfo(f"ERROR","אין שורות המכילות את שם העובד")
                            
                    except pymysql.err.DatabaseError as e:
                        messagebox.showerror("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                        pass

                def Company():
                    try:
                        connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta,  database=database1)
                        cursor = connection.cursor()
                        
                        selected_value = search_ttk.get()

                        query = f"""
                                SELECT compname, workername, COUNT(*) AS total_days_worked 
                                FROM workers 
                                WHERE workername = '{selected_value}' 
                                GROUP BY compname, workername
                            """




                        cursor.execute(query)
                        results = cursor.fetchall()
                        
                        if results:
                            message = "פרטי החברות, מזהה העובד, ומספר הימים שעבד בכל חברה:\n"
                            for result in results:
                                company_name = result[0]
                                workername = result[1]
                                days_worked = result[2]
                                message += f"חברה: {company_name}, מזהה עובד: {workername}, ימים שעבד: {days_worked}\n"
                            messagebox.showinfo("פרטי העבודה בחברות", message)
                        else:
                            messagebox.showinfo("ERROR", "אין נתונים להצגה")
                                        
                    except pymysql.err.DatabaseError as e:
                        messagebox.showerror("אירעה שגיאה בעת התחברות למסד הנתונים:", e)
                        return None

                def hours():
                    try:
                        conn = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                        cursor = conn.cursor()
                        worker_name = search_ttk.get()
                        query = f"SELECT workername, SUM(`hrs`) AS total_wage FROM workers WHERE workername = '{worker_name}' GROUP BY workername"
                        cursor.execute(query)
                        res = cursor.fetchall()
                        total_wage = res[0]
                        if total_wage:
                            messagebox.showinfo('שעות עבודה', f'ס"כ שעות עבודה: {total_wage}')
                        else:
                            messagebox.showinfo('אין תוצאות', 'לא נמצאו נתונים עבור העובד המבוקש')
                    except Exception as e:
                        messagebox.showerror('לא ניתן להתחבר לשרת MySQL', f"לא ניתן היה ליצור חיבור מכיוון שמכונת היעד סירבה לכך באופן פעיל. בדוק את חיבור האינטרנט שלך ונסה שנית. השגיאה היא: {e}")
                    finally:
                        pass

                buttonsframe = cttk.Frame(master=calc_Days_Frame,bootstyle="primary")
                buttonsframe.pack(side=RIGHT,fill='y')



        # ====================== | כפתורי בדיקה |
                def dataAnalysis(text,command):
                    btn2 = ctk.CTkButton(master=buttonsframe,
                                        text=text,
                                        command=command,
                                        corner_radius=0,
                                        border_width=0,
                                        text_color=light,
                                        fg_color=dark,
                                        font=(mainFont,17,'bold'
                                        ))
                    btn2.pack(pady=(0,2))

                dataAnalysis(text='ימי עבודה ',command=howMutchDays)
                dataAnalysis(text='קיבל מפריעה',command=takenMoney)
                dataAnalysis(text='כמה עובדים יש',command=howMutchDaysAll)
                dataAnalysis(text='חברה',command=Company)
                dataAnalysis(text='שעות עבודה',command=hours)
            dataAnalysis_func()
    #----------------------------------------------------------



            # | מס שורה |
            laBel(rowIdFrame,'שורה',light,dark)

    #--------------------------------------------------

            # | שדה מס שורה  |
            iidentry = cttk.Entry(rowIdFrame,background=dark,foreground=danger,font=(mainFont,10,'bold'),textvariable=Row_Number_var,state='disabled',justify='center')
            iidentry.pack(pady=0,side='bottom')
        mainContent()



#============================= | |  tloshim Area   | | ====================================

        f8 = Canvas(self.new_root, bg=dark, bd=0, highlightthickness=0)


        def tloshimSection():
            header = Frame(f8,height=25,bg=primary)
            header.pack(fill='x')

            treeFrame = Frame(header,bg=primary,bd=0,border=0)
            treeFrame.pack(side='top',fill='both')




            t = ttk.Treeview(treeFrame,columns=( 'a','b','c','d','e'))
            t.pack(fill='both')
            t['show']='headings'
            t.heading('e' , text='שם עובד')
            t.heading('d' , text='ת.ז')
            t.heading('c' , text='מתאריך')
            t.heading('b' , text='עד תאריך')
            t.heading('a' , text='תאריך קבלה')



            t.column('a',width=110,)
            t.column('b',width=110,)
            t.column('c',width=110,)
            t.column('d',width=110,)
            t.column('e',width=110,)


            WORKER_NAME_var = StringVar()
            WORKER_ID_var = StringVar()
            FROMS_var = StringVar()
            INTOS_var = StringVar()
            DATE_var = StringVar()

            def get_cursor(event):
                    cursor_row = t.focus()
                    try:    
                        if cursor_row:
                            mainContents = t.item(cursor_row)
                            row = mainContents['values']
                            if row:
                                DATE_var.set(row[0])
                                INTOS_var.set(row[1])
                                FROMS_var.set(row[2])
                                WORKER_ID_var.set(row[3])
                                WORKER_NAME_var.set(row[4])
                                if not row:
                                    print('error')
                    except pymysql.err.IntegrityError as e:
                        print('err', e)
            t.bind("<ButtonRelease-1>", get_cursor)

#===============================================     חיפוש     ================================================
            
            searchframe = Frame(f8,bg=primary,bd=0)
            searchframe.pack(fill='both')
            

            search = ttk.Combobox(searchframe, state='readonly',width=15, font=( mainFont,12 ), justify='center')
            search.pack(side='left', padx=5)
            
    # ---------------- داله البحث في جدول تلوشيم ----------------

            def Search_func():
                try:
                    con = pymysql.connect(host=hostname1, port=porta, user=username1, passwd=passwd1, database=database1)
                    cor = con.cursor()
                    search = search.get()

                    cor.execute(f"SELECT * FROM tloshim WHERE workername = '{search}'")
                    quary = cor.fetchall()

                    if len(quary):
                        t.delete(*t.get_children())
                        for row in quary:
                            t.insert('', END, values=row)
                        t.xview_moveto(1.0)
                        con.commit()
                    else:
                        messagebox.showinfo("החיפוש שלך לא תאם", "לא נמצאו עובדים התואמים לקריטריוני החיפוש שלך")
                except pymysql.Error as e:
                    messagebox.showerror("שגיאה", "אירעה שגיאה בעת בדיקת מסד הנתונים")
                finally:
                    pass



            def get_workers_names():
                try:
                    connection = pymysql.connect(host=hostname1, user=username1, passwd=passwd1, port=porta, database=database1)
                    cursor = connection.cursor()

                    cursor.execute("SELECT DISTINCT workername FROM tloshim ORDER BY workername ASC;")

                    rows = cursor.fetchall()

                    global worker_names
                    worker_names = [row[0] for row in rows]
                
                    search['values'] = worker_names

                except:
                    pass
            get_workers_names()




            #--------------------| שדות כלט חיפוש | כפתור חיפוש| ----------------------
            def kaftorim():

                getdata_button = ctk.CTkButton(master=searchframe,height=30, text="חפש",font=('arial UI',10,'bold'),fg_color='white',text_color='#8BD279', command=Search_func)
                getdata_button.pack(side='left',padx=10,ipadx=10)

                getdata_button.bind('<Enter>', lambda e: getdata_button.configure(fg_color='black',text_color='white'))
                getdata_button.bind('<Leave>', lambda e: getdata_button.configure(fg_color='white',text_color='#8BD279'))

                #-------------------- | שדות כלט חיפוש| ----------------------

                serLabel = ctk.CTkLabel(master=searchframe,fg_color='#8BD279',text_color='black',text='חיפוש',font=('arial',16,'bold'))
                serLabel.pack(side=LEFT,padx=10)
                
                global search2
                Val = [valuee for valuee in trns3.keys()]
                search2 = ttk.Combobox(master=searchframe,background='#8BD279',font=('arial UI',10,'bold'),state='readonly',values=Val)
                search2.pack(side='right',padx=5)


                global Ee1
                Ee1 = Entry(master=searchframe,justify='center',font=('arial UI',10,'bold'),bd=0,highlightbackground='#8BD279',highlightthickness=2,highlightcolor='yellow')
                Ee1.pack(side='right',padx=20)
                
                Ee1.bind('<Enter>', lambda e: Ee1.config(highlightbackground='yellow'))
                Ee1.bind('<Leave>', lambda e: Ee1.config(highlightbackground='#8BD279'))
            kaftorim()
    #-------------------------------------------------------------------------------




    #-----------------------------
            
            frd = Frame(f8,bg=warning,bd=0,highlightbackground='#194D33',highlightthickness=1)
            frd.pack(fill='both')


            ffr = Frame(frd,bg=danger,bd=3,relief='sunken')
            ffr.pack(side='right',ipadx=5,ipady=1800)

    #-------------------------------------------------------------------------------




    #------------------------------- | שדות | ------------------------------------------------------

            def shemot(text):
                #--------------------- שם עובד-----------------------------
                Lb_NAME = Label(ffr,text=text,bg=danger,fg=light,font=(mainFont,10))
                Lb_NAME.pack()

            def shodot(textvariable):
                Er_NAME = ctk.CTkEntry(ffr,border_width=0,justify='center',placeholder_text_color=light,textvariable=textvariable)
                Er_NAME.pack()

            shemot(text='שם עובד')

            shodot(textvariable=WORKER_NAME_var)

            shemot(text='ת.ז')

            shodot(textvariable=WORKER_ID_var)

            shemot(text='מתאריך')

            shodot(textvariable=FROMS_var)

            shemot(text='עד תאריך')

            shodot(textvariable=INTOS_var)

            shemot(text='תאריך קבלה')

            shodot(textvariable=DATE_var)











    #-------------------------------------------------------------------------------
    #
    #
    #
    #-------------- | פונקציית מחיקת שדות | -----------------

            def Entrys_empty():
                WORKER_NAME_var.set('')
                WORKER_ID_var.set('')
                FROMS_var.set('')
                INTOS_var.set('')
                DATE_var.set('')

    #-------------------------------------------------------------------------------
    #
    #
    #
    # ---------------------  3 داله الاتصال مع قاعدة البيانات صفحه تلوشيم ------------------------------

            def coco():
                try:    
                    con = pymysql.connect(host=hostname1,port=porta, user=username1, passwd=passwd1, database=database1)
                    cor2 = con.cursor()
                    cor2.execute("INSERT INTO tloshim VALUES(%s,%s,%s,%s,%s)", (
                            DATE_var.get(),
                            INTOS_var.get(),
                            FROMS_var.get(),
                            WORKER_ID_var.get(),
                            WORKER_NAME_var.get(),
                                    ))
                    con.commit()
                    con.close()
                    Entrys_empty()
                    fetch_all2()
                except pymysql.err.IntegrityError as e:
                    messagebox.showerror("Error:", e)
                        
                
                    #--------------------- 3  داله جلب البيانات من قاعدة البيانات صفحه تلوشيم ------------------------------


    #-------------- | פונקציית קבלת נתונים | -----------------

            def fetch_all2():
                try:
                    con = pymysql.connect(host=hostname1,port=porta, user=username1, passwd=passwd1, database=database1)
                    cor = con.cursor()
                    cor.execute('SELECT * FROM tloshim')
                    rows = cor.fetchall()
                    if len(rows) != 0:
                        t.delete(*t.get_children())
                        for row in rows:
                            t.insert('', END, values=row)
                    con.commit()
                    con.close()
                except pymysql.err.IntegrityError as e:
                    print("Error:", e)
            fetch_all2()        

    #----------------------------------------------------------------------------------
    #
    #
    #
    #--------------- | כפתור הוספה | ---------------



            Add_tloshim = ctk.CTkButton(frd, text="הוספה",font=('arial UI',17,'bold'),fg_color='orange',text_color='black',width=180,height=60,command=coco)
            Add_tloshim.pack(side='bottom')

            Add_tloshim.bind('<Enter>', lambda e: Add_tloshim.configure(fg_color='black',text_color='white'))
            Add_tloshim.bind('<Leave>', lambda e: Add_tloshim.configure(fg_color='orange',text_color='black'))

    #----------------------------------------------------------------------------------
    #
    #
    #
    #-------------- | Exel פונקציית הדפס לקובץ | -----------------

            def print_treeview():
                try:
                    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")])
                    if file_path:
                        workbook = openpyxl.Workbook()
                        sheet = workbook.active
                        sheet.sheet_view.rightToLeft = True

                        columns = t["columns"]
                        column_headers = [t.heading(column)["text"] for column in columns]

                        column_headers.reverse()

                        sheet.append(column_headers)

                        items = t.get_children()

                        for row_idx, item in enumerate(items, start=2):
                            values = [t.item(item, "values")[columns.index(column)] for column in columns]

                            values.reverse()

                            for col_idx, value in enumerate(values, start=1):
                                sheet.cell(row=row_idx, column=col_idx, value=value)

                        for col_idx, column_header in enumerate(column_headers, start=1):
                            sheet.cell(row=1, column=col_idx, value=column_header)
                            sheet.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = len(column_header) + 2

                        workbook.save(file_path)
                        if workbook:
                            messagebox.showinfo('הצלחה', f"הקובץ נשמר ב: {file_path}")
                except:
                    messagebox.showerror('Error', 'The file should not be open.')

            print_button = ctk.CTkButton(frd,hover_color='#09795f',fg_color='green',height=60,text_color='white',text='Exel לקובץ הדפס ',command=print_treeview)
            print_button.pack(side='left')
    #----------------------------------------------------------------------------------  
    #
    #
    #
    #-------------- | Docx פונקציית הדפס לקובץ | -----------------


            def print_to_word():
                file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
                if file_path:
                    document = Document()
                    
                    title = document.add_heading('מידע על העובד\n', level=1)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                    table_data = []
                    column_names = [t.heading(column)["text"] for column in t["columns"]]

                    table_data.append(column_names)
                        
                    for item in t.get_children():
                        item_values = [t.item(item, "values")[i] for i in range(len(t["columns"]))]

                        table_data.append(item_values)
                        
                    table = document.add_table(rows=len(table_data), cols=len(column_names))
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.cell(row_idx, col_idx)
                            cell.text = str(cell_data)
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                    document.save(file_path)
                    messagebox.showinfo('הפעולה הושלמה בהצלחה', f"{file_path} הקובץ נשמר ב: ")
                    fetch_all2()

    #------------------ | כפתור הדפסה docx | ---------------------
                    
            word_button = ctk.CTkButton(frd,hover_color='#09795f',width=150,height=60,fg_color='#2db4b3',text_color='white', text='Docx לקובץ הדפס ', command=print_to_word)
            word_button.pack(side='right')

    #----------------------------------------------------------------------------------
    #
    #
    #
    #-------------- | כפתור רענון | -----------------
            
            refresh = ctk.CTkButton(frd,font=('arial UI',12,'bold'),hover_color='#141414',text_color='white',text='רענון טבלה',width=30,height=25,fg_color='#0eae88',cursor='hand1', command=fetch_all2)
            refresh.pack(side='top')




    #----------------------------------------------------------------------------------
    #
            def update():
                try:
                    connect = pymysql.connect(host=hostname1,port=porta, user=username1, passwd=passwd1, database=database1)
                    cor = connect.cursor()
                    cor.execute(" UPDATE tloshim SET `takendate`=%s, `todate`=%s, `fromdate`=%s , `id`=%s  WHERE  `workername`=%s", (
                        DATE_var.get(),
                        INTOS_var.get(),
                        FROMS_var.get(),
                        WORKER_ID_var.get(),
                        WORKER_NAME_var.get()
                        ))
                    connect.commit()
                    if connect != 0:
                        messagebox.showinfo('בוצע', 'ערוך בהצלחה')
                        fetch_all2()
                except pymysql.Error as e:
                    messagebox.showerror(f"אירעה שגיאה: {e}")
                finally:
                    if connect:
                        connect.close()

            refresh = ctk.CTkButton(frd,font=('arial UI',20,'bold'),hover_color='red',text_color='white',text='עדכון',width=30,height=25,fg_color='#0eae88',cursor='hand1', command=update)
            refresh.pack(side='top')
        tloshimSection()


        self.new_root.mainloop()



if __name__ == '__main__':
    MainScreen()
